from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from PIL import Image  # Make sure to install the Pillow library

def get_image_dimensions(image_path):
    # Open the image file and get its original dimensions
    with Image.open(image_path) as img:
        width, height = img.size
    return width, height

def insert_image_in_footer(doc, image_path):
    # Iterate through sections in the document
    for section in doc.sections:
        # Access the footer of the section
        footer = section.footer

        # Create a new paragraph in the footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()

        # Set paragraph alignment to center (adjust as needed)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Get the original dimensions of the image
        width, height = get_image_dimensions(image_path)

        # Insert the image into the paragraph with original size
        run = paragraph.add_run()
        run.add_picture(image_path, width=Pt(width), height=Pt(height))

        # Add a break to move to the next line (optional)
        paragraph.add_run().add_break()

def main():
    # Open an existing DOCX file
    doc_path = '../01. Schedule-Masking-NoFooter.docx'
    doc = Document(doc_path)

    # Insert an image into the footer of each section with original size
    image_path = '../footer.png'
    insert_image_in_footer(doc, image_path)

    # Save the modified document
    output_path = 'output2.docx'
    doc.save(output_path)

if __name__ == "__main__":
    main()
