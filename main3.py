from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def insert_image_in_footer(doc, image_path, width=None, height=None):
    # Iterate through sections in the document
    for section in doc.sections:
        # Access the footer of the section
        footer = section.footer

        # Create a new paragraph in the footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()

        # Set paragraph alignment to center (adjust as needed)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Insert the image into the paragraph with specified width and height
        run = paragraph.add_run()
        run.add_picture(image_path, width=width)

        # Add a break to move to the next line (optional)
        paragraph.add_run().add_break()

def add_footer(input_path, output_doc_path, footer_path):
    # Open an existing DOCX file
    doc_path = input_path
    doc = Document(doc_path)

    # Insert an image into the footer of each section with specified size
    image_path = footer_path
    width = Inches(8.237)  # Specify the desired width in inches
    height = Inches(1.056)  # Specify the desired height in inches
    insert_image_in_footer(doc, image_path, width, height)

    # Save the modified document
    output_path = output_doc_path
    doc.save(output_path)

def main():
    add_footer('01. Schedule-Masking-NoFooter.docx', "01. Schedule-Masking-WithFooter.docx", "footer.png")
    add_footer('03. Wording-Masking-NoFooter.docx', "03. Wording-Masking-WithFooter.docx", "footer.png")

if __name__ == "__main__":
    main()
