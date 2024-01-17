# from docx import Document
# from docx.shared import Inches

# def insert_image(doc, image_path):
#     doc.add_picture(image_path, width=Inches(1.0))  # Adjust width as needed

# def main():
#     # Open an existing DOCX file
#     doc_path = '../01. Schedule-Masking-NoFooter.docx'
#     doc = Document(doc_path)

#     # Insert an image into the document
#     image_path = '../footer.png'
#     insert_image(doc, image_path)

#     # Save the modified document
#     output_path = 'output.docx'
#     doc.save(output_path)

# if __name__ == "__main__":
#     main()

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement

def insert_image_in_footer(doc, image_path):
    # Iterate through sections in the document
    for section in doc.sections:
        # Access the footer of the section
        footer = section.footer

        # Create a new paragraph in the footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()

        # Set paragraph alignment to center (adjust as needed)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Insert the image into the paragraph
        run = paragraph.add_run()
        run.add_picture(image_path, width=Inches(1.0))  # Adjust width as needed

        # Add a break to move to the next line (optional)
        paragraph.add_run().add_break()

def main():
    # Open an existing DOCX file
    doc_path = '../01. Schedule-Masking-NoFooter.docx'
    doc = Document(doc_path)

    # Insert an image into the footer of each section
    image_path = '../footer.png'
    insert_image_in_footer(doc, image_path)

    # Save the modified document
    output_path = 'output.docx'
    doc.save(output_path)

if __name__ == "__main__":
    main()
